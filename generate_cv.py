from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

styles = doc.styles

# Helper: set paragraph spacing
def set_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line)

# Helper: add run with formatting
def add_run(para, text, bold=False, italic=False, size=10, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

# Helper: section header
def section_header(text):
    para = doc.add_paragraph()
    set_spacing(para, before=10, after=3)
    run = add_run(para, text.upper(), bold=True, size=10, color=(30, 90, 160))
    # Bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E5AA0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

# ── HEADER ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, after=2)
add_run(p, 'Jonathan Sansó', bold=True, size=20, color=(20, 60, 120))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p2, after=2)
add_run(p2, 'Full-Stack / AI Engineer', bold=True, size=12, color=(30, 90, 160))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p3, after=6)
add_run(p3, 'Buenos Aires, Argentina  ·  +54 9 11 6912-3268  ·  jonasans2@live.com.ar', size=9, color=(80, 80, 80))

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p4, after=8)
add_run(p4, 'linkedin.com/in/jonathan-sanso-fullstack  ·  github.com/jonathansansok  ·  portfolio-sanso-jonathan.netlify.app', size=9, color=(30, 90, 160))

# ── PERFIL ──────────────────────────────────────────────────────────────────
section_header('Perfil profesional')

p = doc.add_paragraph()
set_spacing(p, after=2)
add_run(p,
    'Ingeniero Full-Stack / AI con 4 años de experiencia construyendo productos SaaS multi-tenant de principio a fin. '
    'Especializado en plataformas POS, pipelines de reconocimiento facial, automatización de documentos con IA (OCR/NLP/STT), '
    'migraciones de bases de datos legadas (iBase8 → MySQL) y flujos de email transaccional con SendGrid. '
    'Entrego ciclos semanales alineados con CTO y PM, con infraestructura contenerizada y resiliencia zero-downtime.',
    size=9.5)

# ── EXPERIENCIA ──────────────────────────────────────────────────────────────
section_header('Experiencia')

# Ocean Stack
p = doc.add_paragraph()
set_spacing(p, before=6, after=1)
add_run(p, 'Ocean Stack', bold=True, size=10)
add_run(p, '  ·  Full-Stack Engineer', size=10)

p = doc.add_paragraph()
set_spacing(p, after=3)
add_run(p, 'Ene 2024 – presente  ·  Remoto', italic=True, size=9, color=(100, 100, 100))

bullets_niappa = [
    'Construí Niappa POS: plataforma multi-tenant para cadenas de restaurantes. Gestión de menús, stock, proveedores, pedidos con cuentas divididas y pagos Mercado Pago desde un back-office centralizado.',
    'Desarrollé Oceans HR (ATS): pipeline Kanban con drag-and-drop, motor de matching de CVs con IA (Gemini API, cascada 3 modelos, 8 criterios ponderados) y automatización de emails SendGrid por cambio de fase del candidato.',
    'Toda la infraestructura incluye aislamiento por tenant, RBAC, MFA y soporte ES/EN completo.',
    'Stack: Next.js, React, TypeScript, Tailwind, NestJS, Supabase (PostgreSQL + RLS), Redis/BullMQ, Docker, NGINX.',
]
for b in bullets_niappa:
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, after=1)
    add_run(p, b, size=9.5)

# SPF
p = doc.add_paragraph()
set_spacing(p, before=8, after=1)
add_run(p, 'Servicio Penitenciario Federal Argentino', bold=True, size=10)
add_run(p, '  ·  Full-Stack / AI Engineer', size=10)

p = doc.add_paragraph()
set_spacing(p, after=3)
add_run(p, 'Mar 2022 – Dic 2023  ·  Buenos Aires, AR', italic=True, size=9, color=(100, 100, 100))

bullets_spf = [
    'Lideré un ecosistema de dos plataformas: sistema de operaciones internas y portal público de verificación de documentos.',
    'Construí Python Face Matcher en producción: InsightFace buffalo_l, embeddings 512-d, búsqueda coseno con NumPy, almacenamiento BLOB en MariaDB.',
    'Desarrollé pipeline de inteligencia documental con IA: Tesseract OCR, NLP, Speech-to-Text con revisión humana en el loop.',
    'Migré ~110 GB de datos legados iBase8 (tablas relacionales + PDFs, imágenes escaneadas, ZIPs, documentos Word) a MySQL/MariaDB en 80+ módulos.',
    'Implementé verificación segura de documentos: HMAC-SHA256, JWT, marcas de agua, bloqueo de captura de pantalla.',
    'Reemplazé reportes manuales con generación automática de PDF/Excel/Word. Infraestructura: Redis, queues asíncronas, Docker, NGINX, PM2, Debian.',
]
for b in bullets_spf:
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, after=1)
    add_run(p, b, size=9.5)

# VirtuaState
p = doc.add_paragraph()
set_spacing(p, before=8, after=1)
add_run(p, 'VirtuaState', bold=True, size=10)
add_run(p, '  ·  Frontend Developer', size=10)

p = doc.add_paragraph()
set_spacing(p, after=3)
add_run(p, 'May 2022 – Dic 2023  ·  Remoto', italic=True, size=9, color=(100, 100, 100))

bullets_vs = [
    'Construí sitio de marketing para estudio VR/AR (www.virtuastate.net): HTML5, CSS, JavaScript, diseño responsive mobile-first.',
    'Automaticé pipeline de assets 3D con Python (Blender bpy) para exportar escenas a .glb; integración web con <model-viewer> y React Three Fiber.',
    'Migración hacia arquitectura React.js al crecer el scope. SEO, UI/UX alineada con la marca.',
]
for b in bullets_vs:
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, after=1)
    add_run(p, b, size=9.5)

# ── HABILIDADES TÉCNICAS ──────────────────────────────────────────────────────
section_header('Habilidades técnicas')

skills_table_data = [
    ('Frontend',        'React · Next.js · TypeScript · Tailwind CSS · SCSS · Redux · i18n'),
    ('Backend',         'Node.js · NestJS · Express · Python · FastAPI · REST APIs · Prisma · OpenAPI'),
    ('Bases de datos',  'PostgreSQL · Supabase (RLS/Triggers/RPC) · MySQL · MariaDB · MongoDB · Redis · BullMQ'),
    ('IA / ML',         'InsightFace (face recognition) · Tesseract OCR · NLP · STT · Gemini API (LLM cascade) · RAG · Vector Search'),
    ('Email',           'SendGrid (transaccional, automatizado, semi-automatizado, audit log)'),
    ('Infraestructura', 'Docker · NGINX · Debian · PM2 · AWS (EC2, RDS, S3, CloudFront) · Terraform (IaC) · CI/CD · GitHub Actions'),
    ('Seguridad',       'JWT · MFA · HMAC-SHA256 · CSRF · RBAC · Row Level Security · Tenant isolation'),
    ('Herramientas',    'Git · Figma · Agile/Scrum · Jest · Storybook'),
]

table = doc.add_table(rows=len(skills_table_data), cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, (cat, vals) in enumerate(skills_table_data):
    row = table.rows[i]
    c0 = row.cells[0]
    c1 = row.cells[1]
    c0.width = Cm(3.5)
    c1.width = Cm(14)
    p0 = c0.paragraphs[0]
    p1 = c1.paragraphs[0]
    add_run(p0, cat, bold=True, size=9)
    add_run(p1, vals, size=9)
    for cell in (c0, c1):
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)

# ── PROYECTOS DESTACADOS ──────────────────────────────────────────────────────
section_header('Proyectos destacados')

projects = [
    (
        'Niappa POS (Ocean Stack)',
        'SaaS multi-tenant para cadenas de restaurantes. Menús, stock, proveedores, pedidos con cuentas divididas, '
        'reportes PDF/Excel diarios, pagos Mercado Pago, soporte ES/EN. Deploy: niappa-restaurant.vercel.app',
    ),
    (
        'Oceans HR – ATS (Ocean Stack)',
        'Sistema de seguimiento de candidatos con pipeline Kanban, motor de matching IA (Gemini, 8 criterios ponderados), '
        'emails transaccionales SendGrid por fase, RBAC multi-tenant, i18n ES/EN. Deploy: oceans-hr.vercel.app',
    ),
    (
        'Portal Público de Verificación de Documentos [Privado]',
        'Verificación de autenticidad de documentos sin login: HMAC-SHA256, JWT, overlay con marca de agua, '
        'bloqueo de captura. Python Face Matcher (InsightFace buffalo_l) + pipeline OCR/NLP/STT integrados.',
    ),
    (
        'Plataforma para Eventos Complejos V2 [Privada]',
        'Sistema nacional que reemplazó planillas Excel: formularios guiados, mapa interactivo, reportes automáticos '
        'PDF/Excel/Word, alertas WhatsApp en tiempo real, MFA, backups automáticos. Docker + NGINX + Debian.',
    ),
    (
        'Migración iBase8 → MySQL/MariaDB',
        '~110 GB de datos legados (tablas, PDFs, imágenes escaneadas, ZIPs, documentos Word) migrados a MySQL/MariaDB '
        'en 80+ módulos con scripts ETL Python.',
    ),
]

for title, desc in projects:
    p = doc.add_paragraph()
    set_spacing(p, before=5, after=1)
    add_run(p, title, bold=True, size=9.5)
    p2 = doc.add_paragraph()
    set_spacing(p2, after=3)
    add_run(p2, desc, size=9.5)

# ── EDUCACIÓN ──────────────────────────────────────────────────────────────────
section_header('Educación')

p = doc.add_paragraph()
set_spacing(p, before=4, after=1)
add_run(p, 'Tecnicatura Superior en Programación', bold=True, size=9.5)

p = doc.add_paragraph()
set_spacing(p, after=1)
add_run(p, 'Instituto Técnico Superior  ·  Buenos Aires, AR  ·  2021 – 2023', size=9, italic=True, color=(100, 100, 100))

p = doc.add_paragraph()
set_spacing(p, before=4, after=1)
add_run(p, 'Formación complementaria', bold=True, size=9.5)

certs = [
    'CoderHouse – Desarrollo Web, React.js, Backend Node.js',
    'Udemy – Python, Docker, AWS, TypeScript avanzado',
    'Aprendizaje continuo: arquitectura de sistemas, IA/ML aplicada, seguridad web',
]
for c in certs:
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, after=1)
    add_run(p, c, size=9.5)

# ── IDIOMAS ──────────────────────────────────────────────────────────────────
section_header('Idiomas')

p = doc.add_paragraph()
set_spacing(p, before=4, after=4)
add_run(p, 'Español', bold=True, size=9.5)
add_run(p, ' — Nativo     ', size=9.5)
add_run(p, 'Inglés', bold=True, size=9.5)
add_run(p, ' — Avanzado (lectura técnica fluida, comunicación escrita profesional)', size=9.5)

output = r'C:\repos\portf\portfolio-j-sanso\assets\Jonathan-Sanso-CV-ES.docx'
doc.save(output)
print(f'Saved: {output}')
